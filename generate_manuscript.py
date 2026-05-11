#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate manuscript in docx format following Revista Facultad de Ingeniería guidelines.
Then convert to ODT using LibreOffice.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Helper to set double spacing via XML ────────────────────────
def set_double_spacing(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '480')
    spacing.set(qn('w:lineRule'), 'auto')
    # Remove existing spacing elements
    for s in pPr.findall(qn('w:spacing')):
        pPr.remove(s)
    pPr.append(spacing)

def set_style_double_spacing(style):
    pPr = style.paragraph_format._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '480')
    spacing.set(qn('w:lineRule'), 'auto')
    for s in pPr.findall(qn('w:spacing')):
        pPr.remove(s)
    pPr.append(spacing)

# ── Document setup ──────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.59)   # Letter width
section.page_height = Cm(27.94)  # Letter height
section.left_margin = Cm(3)
section.right_margin = Cm(3)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)

# Default Normal style: Arial 12pt, double-spaced, no extra spacing
normal = doc.styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(12)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
set_style_double_spacing(normal)

# Heading styles
for lvl, sz in [(1, 12), (2, 12), (3, 12)]:
    h = doc.styles[f'Heading {lvl}']
    h.font.name = 'Arial'
    h.font.size = Pt(sz)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(0)
    set_style_double_spacing(h)

# ── Helper: add paragraph ───────────────────────────────────────
def p(text='', bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
      style='Normal', space_before=0, space_after=0):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.alignment = align
    if space_before:
        para.paragraph_format.space_before = Pt(space_before)
    if space_after:
        para.paragraph_format.space_after = Pt(space_after)
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
    return para

def p_mixed(parts, align=WD_ALIGN_PARAGRAPH.LEFT, style='Normal',
             space_before=0, space_after=0):
    """parts: list of (text, bold, italic) tuples"""
    para = doc.add_paragraph(style=style)
    para.paragraph_format.alignment = align
    if space_before:
        para.paragraph_format.space_before = Pt(space_before)
    if space_after:
        para.paragraph_format.space_after = Pt(space_after)
    for text, bold, italic in parts:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
    return para

def h(text, level=1):
    return doc.add_heading(text, level=level)

def add_table_simple(headers, rows, caption, above=True):
    """Add a simple table with header row and caption above (for tables)."""
    if above:
        p(caption, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, hdr in enumerate(headers):
        hdr_cells[i].text = hdr
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
            row_cells[c_idx].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT)
    p()  # blank after table

def add_figure_placeholder(caption_text, space_before=12):
    """Add a figure placeholder block."""
    p()
    p(caption_text, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER,
      space_before=space_before)
    p()


# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════

# Spanish title
p('Análisis espacial de la susceptibilidad a los movimientos en masa mediante '
  'Proceso de Poisson No Homogéneo: aplicación en Medellín, Colombia',
  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)

# English title
p('Spatial Analysis of Landslide Susceptibility Using a Non-Homogeneous Poisson '
  'Process: Application in Medellín, Colombia',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

p()  # blank line

# Author
p('Edier Aristizábal¹*', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

p()

# Affiliation
p('¹ Universidad Nacional de Colombia, Sede Medellín, Departamento de '
  'Geociencias y Medio Ambiente, Medellín, Colombia',
  align=WD_ALIGN_PARAGRAPH.CENTER)

p()

# Correspondence
p_mixed([('* Autor para correspondencia: ', False, False),
         ('evaristizabalg@unal.edu.co', False, False)],
        align=WD_ALIGN_PARAGRAPH.CENTER)

p()

# ORCID
p('ORCID: https://orcid.org/0000-0002-2648-2197',
  align=WD_ALIGN_PARAGRAPH.CENTER)

p()

# Conflict of interest
p_mixed([('Conflicto de interés: ', True, False),
         ('Los autores declaran que no existen conflictos de interés.', False, False)],
        align=WD_ALIGN_PARAGRAPH.CENTER)

p()

# Funding
p_mixed([('Financiación: ', True, False),
         ('[INDICAR FUENTE DE FINANCIACIÓN O ESCRIBIR: Este trabajo no recibió '
          'financiación específica de agencias de financiamiento del sector público, '
          'comercial o sin fines de lucro.]', False, False)],
        align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# RESUMEN (Spanish Abstract)
# ══════════════════════════════════════════════════════════════════

p('RESUMEN', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

p('Los movimientos en masa representan una de las amenazas naturales más destructivas '
  'en regiones montañosas tropicales densamente urbanizadas. La modelación estadística '
  'de la susceptibilidad ha recurrido históricamente a enfoques de regresión logística '
  'o aprendizaje automático que tratan el inventario como una respuesta binaria de '
  'presencia-ausencia en una cuadrícula regular, ignorando la naturaleza de proceso '
  'puntual de los datos. Este estudio aplica un Proceso de Poisson No Homogéneo (NHPP) '
  'mediante la aproximación INLA a un inventario de 2504 movimientos en masa en '
  'Medellín (Colombia). La log-intensidad se modela como función log-lineal '
  'de cuatro predictores derivados de un Modelo Digital de Elevación de 2 m de '
  'resolución: pendiente, aspecto, indicador de suelos finos y cobertura urbana. '
  'El modelo produce estimaciones identificables y físicamente interpretables: '
  'pendiente β̂₁ = 0,901 (IC₉₅%: [0,856; 0,946]), '
  'aspecto β̂₂ = −0,243 ([−0,283; −0,204]), '
  'cobertura urbana β̂₃ = −0,971 ([−1,154; −0,788]) y '
  'suelos finos β̂₄ = −0,709 ([−0,790; −0,628]), '
  'confirmando el control geomorfológico y litológico de la susceptibilidad. '
  'La capacidad discriminativa del modelo es AUC = 0,762, superior a '
  'modelos de proceso puntual topográficos reportados en la literatura para contextos '
  'similares. El enfoque ofrece una base probabilística rigurosa para la evaluación '
  'cuantitativa de la amenaza por movimientos en masa.')

p()
p_mixed([('Palabras clave: ', True, False),
         ('susceptibilidad; movimientos en masa; proceso puntual espacial; '
          'Proceso de Poisson No Homogéneo; INLA; Medellín; Colombia.', False, False)])

p()

# ── ABSTRACT (English) ─────────────────────────────────────────

p('ABSTRACT', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

p('Landslides are among the most destructive natural hazards in densely urbanized '
  'tropical mountainous regions. Statistical modeling of landslide susceptibility '
  'has historically relied on logistic regression and machine learning approaches '
  'that treat the inventory as a binary presence-absence response on a regular grid, '
  'ignoring the point process nature of the data. This study applies a '
  'Non-Homogeneous Poisson Process (NHPP) via the Integrated Nested Laplace '
  'Approximation (INLA) to an inventory of 2,504 landslides in Medellín, Colombia. '
  'The log-intensity is modeled as a log-linear function of four predictors derived '
  'from a 2-m resolution LiDAR Digital Elevation Model: slope, aspect, fine soils '
  'indicator, and urban land cover. '
  'The model produces identifiable and physically interpretable estimates: '
  'slope β̂₁ = 0.901 (95% CI: [0.856; 0.946]), '
  'aspect β̂₂ = −0.243 ([−0.283; −0.204]), '
  'urban cover β̂₃ = −0.971 ([−1.154; −0.788]), and '
  'fine soils β̂₄ = −0.709 ([−0.790; −0.628]), '
  'confirming the geomorphological and lithological control of susceptibility. '
  'Slope emerges as the dominant positive predictor: a one-standard-deviation '
  'increase multiplies the local intensity by e^0.901 ≈ 2.5. '
  'The negative aspect coefficient reflects higher susceptibility on north-facing '
  'slopes where reduced solar radiation maintains elevated soil moisture, increasing '
  'pore-water pressure and reducing shear strength. '
  'The negative coefficients for urban cover and fine soils indicate that, '
  'once topography is controlled, both variables exhibit lower residual susceptibility, '
  'consistent with the preferential location of urban areas on valley floors with '
  'low gradient and the topographic mediation of lithological effects. '
  'The discriminative capacity of the model is AUC = 0.762, which exceeds values '
  'reported for topographic point process models in similar contexts. '
  'Unlike logistic regression or machine learning, the NHPP framework avoids '
  'arbitrary grid discretization, eliminates pseudo-absence ambiguity, and '
  'produces a continuous susceptibility surface interpretable as expected event '
  'density per unit area. The INLA approximation provides Bayesian credibility '
  'intervals for all coefficients with explicit uncertainty quantification. '
  'The proposed NHPP-INLA framework constitutes a rigorous, reproducible, and '
  'interpretable probabilistic alternative for quantitative landslide hazard '
  'assessment in topographically complex terrain.')

p()
p_mixed([('Keywords: ', True, False),
         ('landslide susceptibility; spatial point process; Non-Homogeneous Poisson '
          'Process; Bayesian inference; INLA; Medellín; Colombia.', False, False)])

p()

# ══════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════
h('1. Introducción', level=1)

p('La modelación de la susceptibilidad por movimientos en masa ha evolucionado '
  'desde métodos heurísticos basados en criterios de expertos hasta modelos '
  'estadísticos y de aprendizaje automático, donde la regresión logística es el '
  'enfoque más utilizado [1]. La mayoría de estos métodos estadísticos proyectan '
  'el inventario de eventos sobre una cuadrícula regular de celdas, asignan '
  'etiquetas binarias de presencia-ausencia de movimientos en masa y modelan la '
  'probabilidad de ocurrencia en función de atributos como la pendiente, la '
  'curvatura, la litología o el uso del suelo, entre otros [2, 3]. Aunque '
  'operativamente conveniente, esta discretización conlleva pérdidas de información '
  'espacial importantes y confunde la ausencia verdadera de ocurrencia de movimientos '
  'en masa con la falta de observación o registro en el inventario, lo que sesga la '
  'evaluación de la susceptibilidad [3].')

p()

p('En las últimas décadas, los modelos de aprendizaje automático —bosques aleatorios, '
  'redes neuronales profundas, máquinas de vectores de soporte y sus variantes de '
  'ensamble— han alcanzado niveles de desempeño predictivo excepcionalmente altos en '
  'la evaluación de la susceptibilidad por movimientos en masa, con valores de Área '
  'Bajo la Curva (AUC [0-1]) que frecuentemente superan 0,90 en conjuntos de '
  'validación [1]. Sin embargo, estos modelos operan esencialmente como cajas negras: '
  'su arquitectura interna impide la interpretación directa de las relaciones entre '
  'covariables y respuesta en términos del proceso físico subyacente, dificultando '
  'la validación científica de los resultados y la transferencia del conocimiento '
  'entre contextos geomorfológicos [3]. Esta limitación es particularmente relevante '
  'en el ámbito de los movimientos en masa, donde la identificación cuantitativa de '
  'los factores de control —topografía, litología, uso del suelo— tiene implicaciones '
  'directas para la gestión del riesgo y el diseño de medidas de reducción de la '
  'amenaza.')

p()

p('Los modelos de la familia de los Modelos Lineales Generalizados (GLM), en cambio, '
  'vinculan cada predictor a la respuesta mediante coeficientes interpretables que '
  'cuantifican el efecto de cada variable controlando el resto, permitiendo conectar '
  'los resultados estadísticos con el conocimiento geomorfológico del proceso e '
  'incorporar, bajo un marco bayesiano, información a priori físicamente fundamentada '
  'con cuantificación explícita de la incertidumbre en todas las estimaciones. En '
  'particular, el GLM de Poisson modela el número esperado de eventos en cada unidad '
  'espacial de análisis (celda, unidad de ladera, cuenca, etc.) como función '
  'log-lineal de las covariables, lo que constituye la forma discreta, denominada '
  'modelo espacial de área, de un mecanismo generador más fundamental. En el modelo '
  'de área, la intensidad es constante dentro de cada polígono y el logaritmo del '
  'área se incluye como offset en el predictor lineal para modelar tasas de ocurrencia '
  'independientemente del tamaño de la unidad; la autocorrelación espacial entre '
  'unidades contiguas puede incorporarse mediante componentes autorregresivos '
  'condicionales (CAR) [4, 5], dando lugar a modelos lineales generalizados '
  'jerárquicos de Poisson con estructura espacial. Este paradigma opera en el espacio '
  'discreto: la susceptibilidad se expresa como función constante por partes que '
  'cambia abruptamente en los límites de los polígonos, y los resultados pueden '
  'depender del tamaño y la forma de las unidades elegidas, el Problema de la Unidad '
  'de Área Modificable (MAUP; [6, 7]).')

p()

p('Cuando, en cambio, el inventario se concibe como una nube de puntos en el espacio '
  'continuo ℝ², esta misma estructura conduce de manera natural al marco de los '
  'procesos puntuales espaciales, donde la función de intensidad λ(s) varía de forma '
  'continua y suave en todo el dominio sin imponer unidades de agregación predefinidas. '
  'Desde esta perspectiva, el inventario de movimientos en masa es mejor entendido '
  'como una realización de un conjunto aleatorio de puntos en el dominio continuo '
  'Ω ⊂ ℝ², y la inferencia estadística apunta a la función de intensidad '
  'espacialmente variable λ(s) que gobierna el número esperado de eventos por unidad '
  'de área [8]. El modelo más parsimonioso de esta familia es el Proceso de Poisson '
  'No Homogéneo (NHPP): los eventos son condicionalmente independientes dado λ(s), y '
  'esta intensidad se modela como función log-lineal de las covariables espaciales, '
  'la misma estructura del GLM de Poisson, pero formulada en el dominio continuo sin '
  'necesidad de imponer una cuadrícula. Este enfoque evita la discretización '
  'arbitraria en celdas, suprime la ambigüedad de las ausencias (celdas sin '
  'movimientos en masa) y produce una superficie de susceptibilidad continua '
  'directamente interpretable como densidad esperada de eventos por unidad de área.')

p()

p('La aplicación del proceso puntual espacial a los movimientos en masa sigue siendo '
  'escasa en la literatura. Lombardo et al. [9] demostraron el enfoque para flujos '
  'de detritos en Sicilia, y Lombardo et al. [10] lo extendieron a contextos '
  'espacio-temporales. En el contexto andino, Aristizábal et al. [11] han aplicado '
  'recientemente el modelo de Poisson espacial de área, con componente CAR sobre '
  'unidades de ladera, al departamento de Antioquia, Colombia; ese enfoque opera en '
  'el espacio discreto de los polígonos. En este trabajo se implementa un proceso '
  'puntual espacial en el dominio continuo, Proceso de Poisson No Homogéneo (NHPP), '
  'para los movimientos en masa en la ciudad de Medellín, Colombia, con dos '
  'predictores topográficos y considerando la cobertura urbana/rural y los materiales '
  'que conforman las laderas como suelos gruesos o finos.')

# ══════════════════════════════════════════════════════════════════
# 2. ÁREA DE ESTUDIO
# ══════════════════════════════════════════════════════════════════
h('2. Área de estudio', level=1)

p('El Valle de Aburrá, donde se encuentra la ciudad de Medellín, es un valle '
  'intramontano de aproximadamente 1165 km² situado en la Cordillera Central de los '
  'Andes colombianos (6°6\'-6°41\' N, 75°45\'-75°21\' O). La topografía es '
  'característica de los Andes tropicales: las elevaciones oscilan entre '
  'aproximadamente 1300 m s.n.m. en el fondo del valle y más de 2600 m en las '
  'crestas, con flancos empinados disectados por numerosas quebradas tributarias del '
  'río Medellín [12]. Geomorfológicamente, el valle está modelado por la convergencia '
  'oblicua de las placas de Nazca y del Caribe con la placa Sudamericana [13, 14], '
  'que genera un levantamiento activo, fallas y una topografía altamente montañosa '
  'que favorece intrínsecamente los procesos de remoción en masa [15]. La región '
  'alberga más de 3,6 millones de habitantes, de los cuales el 95% reside en áreas '
  'urbanas que ocupan aproximadamente 340 km²; su crecimiento demográfico fue 30 '
  'veces mayor entre 1905 y 2005 [12]. Los movimientos en masa representan la amenaza '
  'natural más recurrente: causan tres de cada diez desastres reportados y el 75% de '
  'las víctimas anuales por eventos naturales en la región [16].')

p()

p('El clima es tropical bimodal, regulado principalmente por la migración meridional '
  'de la Zona de Convergencia Intertropical (ZCIT) [17, 12, 18]. Las precipitaciones '
  'medias anuales oscilan entre 1500 y 2500 mm, con dos temporadas lluviosas bien '
  'definidas: abril-junio (AMJ) y septiembre-noviembre (SON) [17, 12]. La '
  'variabilidad interanual está fuertemente modulada por el fenómeno El Niño-'
  'Oscilación del Sur (ENOS): las fases La Niña incrementan significativamente las '
  'precipitaciones —en particular durante la segunda temporada lluviosa— aumentando '
  'la humedad del suelo antes del inicio de las épocas húmedas, lo que predispone '
  'las laderas a la falla; las fases El Niño las reducen [12, 18]. La lluvia '
  'desencadena el 92% de todos los movimientos en masa registrados en los Andes '
  'colombianos entre 1970 y 2023 [19].')

p()

p('Los terrenos están conformados predominantemente por rocas metamórficas —gneis, '
  'esquistos y anfibolitas— que conforman el basamento de la Cordillera Central '
  '(Complejo Cajamarca), cuyo protolito varía desde sedimentos paleozoicos hasta '
  'arcos volcánicos mesozoicos [14]. Localmente se presentan cuerpos intrusivos '
  'granodioríticos y gabros del Cretácico tardío y el Terciario temprano, rocas '
  'volcánicas andesíticas a basálticas intercaladas con depósitos volcanoclásticos, '
  'y rocas ultramáficas (serpentinitas y peridotitas) producto de secuencias '
  'ofiolíticas [14]. Sobre estas unidades se desarrollan perfiles de meteorización '
  'profundos y heterogéneos: en granitoides, los suelos residuales y saprolitos de '
  'color amarillo-rojizo (Munsell 10YR 7/4) superan los 50 m de espesor, sobre más '
  'de 100 m de roca químicamente desintegrada; en rocas ultrabásicas y metamórficas, '
  'los perfiles son más delgados (10-20 m), de color naranja-rojizo (Munsell 7.5YR '
  '7/6) y con abundantes fragmentos muy meteorizados [20, 15]. Los materiales de '
  'ladera están dominados por suelos residuales sobre roca profundamente meteorizada '
  'y por depósitos coluviales derivados de procesos de remoción en masa de largo '
  'plazo [20].')

# ══════════════════════════════════════════════════════════════════
# 3. DATOS
# ══════════════════════════════════════════════════════════════════
h('3. Datos', level=1)

p('El inventario de movimientos en masa utilizado comprende 2504 eventos '
  '(deslizamientos superficiales, flujos de detritos y deslizamientos translacionales) '
  'compilado mediante interpretación sistemática de imágenes de satélite de resolución '
  'submétrica disponibles en Google Earth entre 2001 y 2024 (Figura 1). A diferencia '
  'de los inventarios basados en registros históricos institucionales, que introducen '
  'un sesgo de detección hacia eventos de mayor magnitud o localizados en zonas '
  'urbanizadas con mayor cobertura mediática, el inventario de percepción remota '
  'permite detectar movimientos en masa de forma espacialmente uniforme a lo largo '
  'del período de disponibilidad de imágenes de alta resolución. Esta estrategia de '
  'compilación reduce el sesgo espacial de observación y produce un inventario más '
  'representativo de la distribución de la susceptibilidad. Los eventos están '
  'georeferenciados como localizaciones puntuales de la zona de iniciación. La '
  'densidad media es de aproximadamente 7,08 eventos km⁻², con agrupamiento alto en '
  'las laderas del valle.')

add_figure_placeholder(
    'Figura 1. Localización espacial de los 2504 movimientos en masa registrados '
    'en el dominio de estudio del Valle de Aburrá. Cada punto rojo representa la '
    'zona de iniciación de un evento. El contorno negro delimita el dominio de '
    'estudio Ω (335,21 km²).')

p('Se usaron covariables topográficas derivadas de un Modelo Digital de Elevación '
  '(MDE) de 1 m de resolución obtenido con tecnología LiDAR, remuestreado a 2 m '
  'mediante promediado por agregación para mejorar la relación señal-ruido en las '
  'covariables morfométricas y reducir el costo computacional (Tabla 1). La pendiente '
  '(β, grados) se define como el ángulo de la máxima pendiente descendente; es el '
  'principal impulsor mecánico de la inestabilidad al aumentar la componente tangencial '
  'del esfuerzo gravitacional. El aspecto (α, grados desde el Norte) es la dirección '
  'de la máxima pendiente descendente y controla la exposición solar, la '
  'evapotranspiración y el contenido de humedad del suelo. Las dos covariables '
  'topográficas se estandarizaron a media cero y desviación estándar unitaria sobre '
  'el dominio raster completo antes del ajuste del modelo.')

p()

add_table_simple(
    headers=['Covariable', 'Mín.', 'Q₂₅', 'Mediana', 'Q₇₅', 'Máx.'],
    rows=[
        ['Pendiente (°)', '0,12', '28,37', '36,84', '43,24', '65,13'],
        ['Aspecto (°)',   '0,16', '86,45', '144,35', '203,99', '359,68'],
    ],
    caption='Tabla 1. Estadísticas descriptivas de las covariables topográficas '
            'estandarizadas en las 2504 localizaciones de movimientos en masa '
            '(MDE a 2 m).',
    above=True
)

p('Por ser un territorio fuertemente antropizado se usaron dos tipos de uso del '
  'suelo: urbano y rural. La incorporación de la geología se realizó categorizando '
  'cada unidad geológica en dos clases binarias:')

p()
p_mixed([('Suelos finos: ', True, False),
         ('incluye esquistos, dunitas, rocas volcánicas y depósitos de flujos de '
          'lodos. Esta categoría cubre aproximadamente el 35% del dominio.', False, False)])

p_mixed([('Suelos gruesos: ', True, False),
         ('incluye gneis, anfibolitas, granodioritas, gabros, migmatitas y depósitos '
          'aluviales. Esta categoría cubre el 65% restante del dominio.', False, False)])

# ══════════════════════════════════════════════════════════════════
# 4. METODOLOGÍA
# ══════════════════════════════════════════════════════════════════
h('4. Metodología', level=1)

h('4.1. Proceso de Poisson No Homogéneo (NHPP)', level=2)

p('El inventario de movimientos en masa se trata como una realización de un proceso '
  'puntual espacial en el dominio continuo Ω ⊂ ℝ², cuya distribución está gobernada '
  'por la función de intensidad λ(s). El modelo adoptado es el Proceso de Poisson '
  'No Homogéneo (NHPP), en el que el reto estadístico es estimar λ(s) en todo el '
  'dominio a partir de las 2504 localizaciones observadas, teniendo en cuenta que '
  'la intensidad varía con la topografía, el tipo de suelo y la cobertura del '
  'territorio. Formalmente, el NHPP establece que los eventos son condicionalmente '
  'independientes dado la función de intensidad λ(s), y el conteo en cualquier área '
  'A sigue N(A) ~ Poisson(∫_A λ(s) ds). Con una intensidad log-lineal en '
  'covariables, el modelo se formula como:')

p()

# Equations - centered
p('{Sᵢ} | λ(·) ~ Poisson_NH(λ(s))',
  align=WD_ALIGN_PARAGRAPH.CENTER)
p_mixed([('', False, False)], align=WD_ALIGN_PARAGRAPH.CENTER)

# Equation display with number
eq_para = doc.add_paragraph(style='Normal')
eq_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
eq_para.add_run('   log λ(s) = η(s)').bold = False
# Add tab and equation number
run_num = eq_para.add_run('                                                              (1)')

p()
p('η(s) = β₀ + β₁ · pendiente(s) + β₂ · aspecto(s) + β₃ · geo_finos(s) + β₄ · urb(s)',
  align=WD_ALIGN_PARAGRAPH.CENTER)
p_mixed([('                                                                              (2)', False, False)])

p()

p('donde pendiente y aspecto son las covariables topográficas estandarizadas (media 0, '
  'd.e. 1); geo_finos ∈ {0,1} el indicador de suelos finos; y urb ∈ {0,1} el '
  'indicador de cobertura urbanizada.')

p()

p('Los coeficientes βⱼ actúan sobre la log-intensidad: βⱼ > 0 implica mayor '
  'intensidad esperada al aumentar la covariable j. Un incremento de una desviación '
  'estándar en una covariable continua multiplica la intensidad local por e^βⱼ. Para '
  'covariables binarias, e^βⱼ es la razón de intensidades entre la clase activa y '
  'la de referencia, controlando el resto.')

p()

p('La estimación estadística de los coeficientes βⱼ del NHPP requiere caracterizar '
  'la distribución posterior conjunta de todos los parámetros dado el patrón observado '
  'de deslizamientos. La función de verosimilitud del proceso puntual cuantifica esta '
  'compatibilidad: penaliza configuraciones que predicen alta intensidad donde no '
  'ocurrieron eventos, y baja intensidad donde sí los hay. Su evaluación involucra '
  'la integral de la intensidad ∫_Ω λ(s) ds sobre el dominio continuo, que no tiene '
  'solución analítica cuando la intensidad varía en el espacio según covariables '
  'continuas; esta integral se aproxima numéricamente mediante el dispositivo de '
  'Berman-Turner [8] sobre una malla triangular.')

p()

p('La estimación bayesiana completa de los parámetros requeriría en principio el '
  'muestreo por cadenas de Markov (MCMC), cuyas iteraciones resultan '
  'computacionalmente prohibitivas al evaluar la integral sobre una malla de 92657 '
  'nodos, en este caso, con alta resolución espacial: cada iteración del muestreador '
  'exige recalcular la intensidad en todos los nodos, lo que hace demasiado exigente '
  'computacionalmente el MCMC para mallas de este tamaño en tiempos razonables. Por '
  'ello, se adopta la Aproximación de Laplace Anidada Integrada (INLA; [21, 22]), '
  'que computa las distribuciones marginales posteriores de todos los parámetros '
  'mediante aproximaciones analíticas de Laplace, sin simulación estocástica, con '
  'precisión comparable al MCMC y sin el costo computacional. INLA aplica la '
  'aproximación sobre el predictor lineal con integración numérica sobre una grilla '
  'de configuraciones, proporcionando una cuantificación robusta de la incertidumbre.')

p()

p('La malla triangular cubre tanto el dominio interior como un buffer exterior de '
  '2 km (Figura 2, Tabla 2). La malla es necesaria para aproximar numéricamente la '
  'integral ∫_Ω λ(s) ds mediante el dispositivo de Berman-Turner [8]: los nodos de '
  'la malla actúan como puntos de integración con pesos proporcionales al área de la '
  'celda de Voronoi dual. Las aristas interiores de hasta 100 m proporcionan '
  'suficiente resolución para representar la variabilidad espacial de las covariables '
  'a 2 m. El buffer exterior reduce los artefactos de borde.')

add_figure_placeholder(
    'Figura 2. Malla triangular superpuesta al hillshade del relieve de Medellín. '
    'Puntos rojos: zonas de iniciación de movimientos en masa observados (2504 eventos). '
    'Triangulación azul: malla de integración (92657 nodos, 185024 triángulos). '
    'El buffer exterior (triángulos de mayor tamaño) se extiende 2 km fuera del '
    'dominio de estudio (contorno negro) para atenuar los efectos de borde en la '
    'integración numérica.')

add_table_simple(
    headers=['Parámetro', 'Valor', 'Unidad'],
    rows=[
        ['Longitud máxima de arista (interior)', '100', 'm'],
        ['Longitud máxima de arista (exterior)', '500', 'm'],
        ['Separación mínima entre nodos', '20', 'm'],
        ['Offset interior', '200', 'm'],
        ['Offset exterior', '2000', 'm'],
        ['Número de nodos', '92 657', '—'],
        ['Número de triángulos', '185 024', '—'],
    ],
    caption='Tabla 2. Parámetros de la malla triangular utilizada para la '
            'integración numérica de la verosimilitud del NHPP (dispositivo '
            'de Berman-Turner).',
    above=True
)

p('El modelo se implementó en el lenguaje R [23], utilizando los paquetes INLA [22], '
  'inlabru [24], fmesher, sf y terra. Los coeficientes de regresión se les asignaron '
  'priors débilmente informativos: βⱼ ~ N(0, 1), lo que implica que un incremento '
  'unitario en una covariable se asocia a priori con un factor de intensidad de '
  'e^±1 ≈ 2,7 en cualquier dirección. Los mapas predictivos de λ(s) se obtuvieron '
  'sobre grillas de 100 m con 500 muestras de Monte Carlo de la posterior conjunta. '
  'El desempeño discriminativo se evaluó con el AUC calculado sobre una grilla '
  'hexagonal de ≈ 0,25 km², asignando a cada celda un indicador binario (1 si '
  'contenía al menos un evento) y como puntuación la intensidad total predicha.')

# ══════════════════════════════════════════════════════════════════
# 5. RESULTADOS
# ══════════════════════════════════════════════════════════════════
h('5. Resultados', level=1)

p('La distribución de movimientos en masa de acuerdo con el inventario es '
  'marcadamente no homogénea, con densidades locales que superan los 20 eventos km⁻² '
  'en los flancos de mayor pendiente del valle. La Figura 1 muestra la distribución '
  'espacial de los eventos y la Figura 2 la malla triangular de integración.')

p()

p('De los 2504 eventos, aproximadamente el 35% se localiza en el estrato de suelos '
  'finos y el 65% en suelos gruesos, coherente con las proporciones de área de cada '
  'estrato. Sin embargo, la densidad de eventos por km² es notablemente mayor en '
  'suelos finos. La distribución espacial del inventario refleja fundamentalmente '
  'los patrones topográficos del Valle de Aburrá, sin la concentración en zonas '
  'urbanas que caracteriza a los inventarios históricos; esto es indicativo de que '
  'el inventario de percepción remota captura efectivamente la variabilidad espacial '
  'de la susceptibilidad a escala de ladera.')

p()

p('La Tabla 3 y la Figura 3 presentan los coeficientes posteriores del modelo NHPP. '
  'La pendiente es el único predictor topográfico positivo y significativo; el aspecto '
  'muestra un efecto negativo significativo. Los dos predictores binarios (suelos '
  'finos y cobertura urbana) son negativos y significativos.')

p()

add_table_simple(
    headers=['Predictor', 'Media', 'D.E.', 'IC₂,₅%', 'IC₉₇,₅%'],
    rows=[
        ['Intercepto',           '−11,769', '0,034', '−11,835', '−11,703'],
        ['Pendiente (estand.)',   '0,901',   '0,023', '0,856',   '0,946'],
        ['Aspecto (estand.)',     '−0,243',  '0,020', '−0,283',  '−0,204'],
        ['Cobertura urbana',      '−0,971',  '0,093', '−1,154',  '−0,788'],
        ['Suelos finos (bin.)',   '−0,709',  '0,041', '−0,790',  '−0,628'],
    ],
    caption='Tabla 3. Resumen posterior de los efectos fijos del modelo NHPP '
            '(MDE a 2 m, covariables estandarizadas). Los coeficientes indican '
            'el cambio en log λ(s) por incremento de una desviación estándar en '
            'cada covariable continua, o el efecto de pertenecer a la clase activa '
            'para las variables binarias.',
    above=True
)

add_figure_placeholder(
    'Figura 3. Medias posteriores e intervalos de credibilidad al 95% de los '
    'coeficientes de efectos fijos β̂ (escala log-intensidad) del modelo NHPP '
    '(MDE 2 m). La pendiente es el predictor positivo dominante; aspecto, '
    'cobertura urbana y suelos finos muestran efectos negativos significativos.')

p_mixed([('Pendiente ', True, False),
         ('(β̂₁ = 0,901; IC₉₅%: [0,856; 0,946]): el coeficiente es positivo y '
          'preciso, coherente con el papel mecánico fundamental de la pendiente en '
          'el factor de seguridad. Un incremento de una desviación estándar en la '
          'pendiente multiplica la intensidad por e^0,901 ≈ 2,5.', False, False)])

p()

p_mixed([('Aspecto ', True, False),
         ('(β̂₂ = −0,243; IC₉₅%: [−0,283; −0,204]): el coeficiente negativo indica '
          'mayor susceptibilidad en laderas de orientación norte (aspecto '
          'estandarizado bajo), donde la menor insolación reduce la '
          'evapotranspiración y mantiene mayor humedad del suelo, aumentando las '
          'presiones de poros y reduciendo la resistencia al corte.', False, False)])

p()

p_mixed([('Indicador de suelos finos ', True, False),
         ('(β̂₃ = −0,709; IC₉₅%: [−0,790; −0,628]): coeficiente negativo al parecer '
          'controlado por la topografía. Este resultado sugiere que la mayor '
          'susceptibilidad de los suelos finos está mediada casi completamente por '
          'su topografía asociada (pendientes más elevadas en el estrato fino), y '
          'que una vez que se controlan esas covariables el efecto litológico directo '
          'es débil o de signo contrario al esperado.', False, False)])

p()

p_mixed([('Cobertura urbana ', True, False),
         ('(β̂₄ = −0,971; IC₉₅%: [−1,154; −0,788]): el coeficiente es '
          'significativamente negativo. Los sectores clasificados como urbanos '
          'muestran una intensidad predicha significativamente menor que las zonas '
          'rurales una vez que se controla la topografía. Este resultado refleja que '
          'las zonas urbanas de Medellín se localizan predominantemente en el fondo '
          'plano del valle (baja pendiente) y que el indicador binario captura el '
          'efecto residual de la urbanización, que en pendientes bajas reduce la '
          'susceptibilidad respecto a laderas rurales.', False, False)])

p()

p('Las Figuras 4a y 4b muestran la superficie de intensidad media posterior '
  'λ̂(s) y su incertidumbre en el dominio de estudio. Las áreas de mayor '
  'intensidad (colores rojos) se localizan a lo largo de los flancos del río '
  'Medellín y sus tributarios principales, donde la combinación de alta pendiente '
  'y orientación norte maximiza el predictor lineal. Los colores verdes indican baja '
  'susceptibilidad, correspondientes al fondo plano del valle y a las zonas urbanas. '
  'La distribución espacial de la incertidumbre posterior refleja el número local '
  'de eventos: zonas con alta densidad de deslizamientos exhiben menor varianza '
  'predictiva.')

add_figure_placeholder(
    'Figura 4. Mapas de (a) la intensidad media posterior λ̂(s) (eventos m⁻²) '
    'y (b) la desviación estándar posterior del modelo NHPP en el Valle de Aburrá. '
    '(a) Verde: baja susceptibilidad; rojo: alta susceptibilidad (escala '
    'logarítmica). (b) De violeta a amarillo: baja a alta incertidumbre en la '
    'estimación de λ(s).')

p('El mapa de residuos de Pearson (n_obs − n_esp)/√n_esp (Figura 5) no muestra '
  'un patrón espacial sistemático con estructura geográfica definida, lo que indica '
  'que el modelo no presenta sesgo espacial evidente. Los residuos positivos más '
  'pronunciados se concentran puntualmente en algunas quebradas, posiblemente donde '
  'los registros históricos son incompletos o donde factores antrópicos locales '
  'amplifican la susceptibilidad.')

add_figure_placeholder(
    'Figura 5. Mapa de residuos de Pearson (n_obs − n_esp)/√n_esp del modelo NHPP '
    'sobre el Valle de Aburrá. Rojo: celdas con más eventos de los esperados '
    '(posible subestimación de la susceptibilidad local o subreporte en el '
    'inventario). Azul: celdas con menos eventos de los esperados.')

p('La Tabla 4 resume las métricas de ajuste del modelo. El AUC indica la capacidad '
  'discriminativa del modelo: clasificar correctamente una celda con movimiento en '
  'masa por encima de una celda sin evento en aproximadamente el 76% de los casos. '
  'La Figura 6 muestra la curva ROC completa.')

p()

add_table_simple(
    headers=['Métrica', 'Valor'],
    rows=[
        ['AUC (ROC)',               '0,762'],
        ['Número total de eventos', '2504'],
        ['Intensidad media',        '7,08 eventos km⁻²'],
        ['Área del dominio',        '335,21 km²'],
    ],
    caption='Tabla 4. Métricas de ajuste del modelo NHPP sobre el dominio completo '
            '(MDE 2 m, cuatro predictores).',
    above=True
)

add_figure_placeholder(
    'Figura 6. Curva ROC del modelo NHPP evaluada sobre una grilla hexagonal de '
    '≈ 0,25 km² en el dominio completo. El eje horizontal es la tasa de falsos '
    'positivos y el vertical la tasa de verdaderos positivos. La línea diagonal '
    'punteada representa un clasificador aleatorio (AUC = 0,5).')

# ══════════════════════════════════════════════════════════════════
# 6. DISCUSIÓN
# ══════════════════════════════════════════════════════════════════
h('6. Discusión', level=1)

p('La susceptibilidad a los movimientos en masa puede modelarse mediante '
  'distribuciones de Poisson bajo dos paradigmas espaciales que conviene contrastar. '
  'El modelo de área agrega los eventos en polígonos discretos —celdas regulares, '
  'cuencas o unidades de ladera— y modela el conteo yᵢ en la unidad i como '
  'yᵢ ~ Poisson(λᵢ Aᵢ) [11]. Este enfoque fue aplicado recientemente en el contexto '
  'andino colombiano para el departamento de Antioquia, demostrando su utilidad a '
  'escala regional con unidades de ladera como unidad mínima de análisis [11]. '
  'Sus ventajas son la eficiencia computacional, requiere únicamente una sumatoria '
  'sobre los polígonos, y la interpretabilidad directa de los efectos aleatorios '
  'espaciales como variaciones regionales de la susceptibilidad no explicadas por '
  'los predictores fijos. La principal limitación es que la intensidad estimada es '
  'constante dentro de cada polígono independientemente de la variabilidad '
  'topográfica interna, y que los resultados dependen del tamaño y la forma de las '
  'unidades de análisis elegidas, lo que puede introducir discontinuidades artificiales '
  'en los límites de los polígonos.')

p()

p('El proceso puntual espacial adoptado en este trabajo supera estas limitaciones '
  'estructurales. La función de intensidad λ(s) varía de forma continua en cada '
  'coordenada del dominio, utilizando las covariables exactas en ese punto y '
  'produciendo una superficie de susceptibilidad suave sin discontinuidades en los '
  'límites de polígonos ni dependencia del tamaño de las unidades de análisis. En '
  'los Andes, donde la pendiente puede cambiar en distancias de pocos metros, esta '
  'propiedad es fundamental para representar fielmente la distribución espacial de '
  'la amenaza a escala de ladera. Además, el NHPP no requiere offset de área, la '
  'integral ∫_A λ(s) ds da directamente el número esperado de eventos en cualquier '
  'región A, ni demanda definir celdas sin deslizamientos. La contrapartida es un '
  'mayor costo computacional: la integral ∫_Ω λ(s) ds se evalúa numéricamente sobre '
  'una malla triangular (92657 nodos en este trabajo) y la inferencia requiere la '
  'aproximación bayesiana INLA para ser factible en tiempos razonables. La elección '
  'entre ambos paradigmas depende de la escala de análisis y los objetivos del '
  'estudio: el modelo de área es adecuado a escalas regionales con unidades '
  'geográficas relativamente homogéneas, mientras que el proceso puntual es '
  'preferible cuando se dispone de inventarios de alta resolución y se requiere '
  'estimar la susceptibilidad con resolución métrica en terrenos topográficamente '
  'complejos.')

p()

p('En cuanto a los resultados, los coeficientes del modelo confirman que la '
  'topografía es el principal determinante de la distribución espacial de los '
  'movimientos en masa en Medellín. El coeficiente de pendiente β̂₁ = 0,901 es '
  'positivo y de gran magnitud, reflejando que la aceleración gravitacional tangencial '
  'es el determinante primario de la inestabilidad. Un incremento de una desviación '
  'estándar en pendiente multiplica la intensidad local por e^0,901 ≈ 2,5.')

p()

p('El aspecto negativo (β̂₂ = −0,243) indica mayor susceptibilidad en laderas con '
  'un mayor dirección azimutal, que reciben menos radiación solar directa, mantienen '
  'mayor humedad del suelo durante los períodos de lluvia y desarrollan suelos '
  'potencialmente más profundos por menor evapotranspiración. Este efecto es '
  'independiente de la pendiente y la litología.')

p()

p('El índice topográfico de humedad (TWI = ln[a/tan β]), que integra el área '
  'aportante acumulada aguas arriba y la pendiente local como proxy de la '
  'predisposición de las laderas a la saturación del suelo, fue evaluado como '
  'predictor adicional del modelo NHPP pero resultó estadísticamente no significativo '
  'y fue excluido del modelo final. Cuatro factores pueden explicar este resultado '
  'en el contexto del Valle de Aburrá. En primer lugar, a resoluciones de 2 m, el '
  'cálculo del área aportante es altamente sensible a las microirregularidades del '
  'MDE LiDAR y a los algoritmos de dirección de flujo: en laderas empinadas y de '
  'extensión lateral reducida, dominantes en Medellín, las cuencas de contribución '
  'se reducen a pocos píxeles y la variabilidad numérica del TWI queda dominada por '
  'el ruido topográfico del raster más que por la señal hidrológica real. En segundo '
  'lugar, Medellín es un territorio altamente urbanizado cuya red de drenaje ha sido '
  'profundamente modificada: alcantarillados, cunetas, canalizaciones y obras de '
  'urbanización interceptan y desvían los flujos superficiales de sus trayectorias '
  'topográficas naturales, invalidando el supuesto fundamental del TWI de que el '
  'flujo de agua sigue el gradiente topográfico. En tercer lugar, los movimientos en '
  'masa del inventario corresponden predominantemente a deslizamientos superficiales '
  'y flujos de detritos detonados por lluvias intensas, procesos en los que la '
  'inestabilidad es controlada primariamente por la infiltración vertical, no por la '
  'acumulación de flujo lateral en cuencas de contribución extensas; en este tipo de '
  'procesos, los resultados señalan que el TWI no constituye un predictor físicamente '
  'relevante a escala de ladera. Finalmente, la colinealidad estructural entre el '
  'TWI y la pendiente, componente del denominador de la fórmula, produce inflación '
  'de los intervalos de credibilidad de ambos coeficientes cuando se incluyen '
  'simultáneamente, haciendo que el TWI resulte indistinguible de cero condicionado '
  'en la pendiente ya incorporada al modelo.')

p()

p('El coeficiente negativo del indicador de suelos finos (β̂₃ = −0,709) sugiere que '
  'el diferencial de susceptibilidad entre litologías es capturado casi completamente '
  'por las covariables topográficas: los suelos finos ocupan zonas con pendientes '
  'más altas (que ya elevan la intensidad predicha), y una vez que se controla esa '
  'topografía el efecto litológico directo es estadísticamente negativo o nulo. Este '
  'fenómeno de mediación por covariables topográficas es un hallazgo metodológico '
  'relevante: la litología actúa principalmente como un control de la morfología '
  '(que determina dónde se acumulan los materiales finos), más que como un predictor '
  'independiente de la susceptibilidad condicionada en la topografía.')

p()

p('El coeficiente de cobertura urbana (β̂₄ = −0,971) es fuertemente negativo y '
  'significativo. Los sectores urbanos de Medellín se localizan predominantemente '
  'en el fondo plano del valle y, una vez controlada la topografía, la '
  'impermeabilización de suelos y las obras de estabilización en laderas urbanas '
  'reducen la susceptibilidad residual respecto a laderas rurales.')

p()

p('El AUC de 0,762 es superior al reportado por Lombardo et al. [9] para modelos '
  'de proceso puntual en Sicilia (AUC ≈ 0,63-0,67). Esto es notable considerando '
  'que el modelo NHPP es parsimonioso: basa toda su discriminación en cuatro '
  'predictores fijos, sin componentes no paramétricas adicionales. El resultado '
  'sugiere que, para Medellín, la variabilidad topográfica y litológica explica una '
  'fracción sustancial de la distribución espacial de la amenaza.')

p()

p('El AUC se evalúa a una escala de 500 m (celda hexagonal), que es la escala de '
  'información del modelo. A escalas más gruesas (cuadrantes de 1 km²) el AUC sería '
  'mayor por la reducción de la varianza de Poisson; a escala de punto individual '
  'sería menor por la estocasticidad intrínseca del proceso. La comparación adecuada '
  'con otros enfoques (regresión logística, modelos de aprendizaje automático) '
  'requeriría la misma escala de evaluación y el mismo dominio de integración.')

p()

p('Una ventaja metodológica central de este estudio es el uso de un inventario '
  'basado exclusivamente en percepción remota, construido mediante interpretación '
  'sistemática de imágenes de satélite de alta resolución. Los inventarios históricos '
  'están sujetos a sesgo de detección: registran preferentemente eventos con víctimas '
  'o daños materiales reportados, eventos cercanos a la infraestructura urbana, o '
  'eventos que reciben cobertura mediática [25]. Esta selectividad induce patrones '
  'espaciales sesgados en el inventario que se trasladan directamente al modelo de '
  'susceptibilidad: zonas con alta cobertura de datos históricos parecen más '
  'susceptibles independientemente de sus atributos topográficos. Al construir el '
  'inventario a partir de interpretación visual sistemática de Google Earth para el '
  'período 2001-2024, se logra una cobertura espacialmente más uniforme, condicionada '
  'principalmente a la disponibilidad de imágenes de alta resolución y no al sesgo '
  'del sistema de reporte institucional.')

p()

p('Este enfoque no elimina completamente el sesgo de detección (la cobertura nubosa '
  'y la disponibilidad temporal de imágenes crean variabilidad espacial en la '
  'probabilidad de detección), pero lo reduce sustancialmente respecto a inventarios '
  'históricos. El uso del marco del proceso puntual espacial es especialmente adecuado '
  'para este tipo de inventario: la verosimilitud del NHPP integra la intensidad sobre '
  'todo el dominio sin requerir ausencias explícitas, lo que reduce la sensibilidad '
  'al sesgo de detección uniforme.')

p()

p('El uso de inventarios históricos en este tipo de análisis exige la incorporación '
  'de variables de intervención antrópica (densidad de cortes de talud, distancia a '
  'vías, distancia a rellenos) para capturar el control local del componente antrópico '
  'en la susceptibilidad. Sin embargo, el alcance del actual trabajo se enfoca en la '
  'susceptibilidad asociada a las condiciones naturales del territorio.')

# ══════════════════════════════════════════════════════════════════
# 7. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════
h('7. Conclusiones', level=1)

p('Este estudio implementó un Proceso de Poisson No Homogéneo (NHPP) con inferencia '
  'INLA para modelar la susceptibilidad a los movimientos en masa en Medellín, '
  'Colombia, usando covariables derivadas de un MDE de 2 m. Los resultados confirman '
  'que la topografía es el principal determinante de la distribución espacial de los '
  'movimientos en masa en el territorio. El coeficiente de pendiente '
  'β̂_pendiente = 0,901 (e^0,901 ≈ 2,5×) es el predictor más influyente del modelo, '
  'coherente con el papel mecánico fundamental de la pendiente en el mecanismo del '
  'talud infinito. El aspecto muestra un efecto negativo significativo '
  '(β̂_aspecto = −0,243), indicando mayor susceptibilidad en laderas de orientación '
  'con una dirección azimutal mayor, donde probablemente la menor insolación favorece '
  'la retención de humedad del suelo. El índice topográfico de humedad (TWI), '
  'evaluado como predictor adicional, resultó estadísticamente no significativo y '
  'fue excluido del modelo final, debido a la colinealidad con la pendiente, al '
  'ruido numérico que introduce su cálculo a resoluciones de 2 m en laderas cortas '
  'y empinadas, y a la modificación antrópica de la red de drenaje en un territorio '
  'fuertemente urbanizado.')

p()

p('El indicador de suelos finos muestra un efecto negativo (β̂ = −0,709) una vez '
  'controlada la topografía, evidenciando que la mayor susceptibilidad de los '
  'materiales de grano fino está mediada fundamentalmente por su topografía asociada: '
  'estos materiales ocupan preferentemente zonas de pendiente elevada que ya elevan '
  'la intensidad predicha por el modelo, y el efecto litológico directo resulta '
  'estadísticamente débil. La cobertura urbana es también significativamente negativa '
  '(β̂ = −0,971), reflejando que las áreas urbanizadas se localizan predominantemente '
  'en el fondo plano del valle y que, una vez controlada la topografía, las obras de '
  'estabilización e impermeabilización asociadas a la urbanización reducen la '
  'susceptibilidad residual respecto a laderas rurales.')

p()

p('El modelo NHPP con cuatro predictores fijos proporciona estimaciones robustas e '
  'interpretables con AUC = 0,762, superior al reportado para modelos de proceso '
  'puntual topográficos en contextos similares en la literatura. El enfoque NHPP-INLA '
  'evita las decisiones arbitrarias inherentes a la discretización en celdas '
  '—resolución de la grilla, selección de ausencias— y produce intervalos de '
  'credibilidad bayesianos sobre todos los coeficientes con cuantificación explícita '
  'de la incertidumbre. Este marco constituye una alternativa probabilística rigurosa '
  'y reproducible a la regresión logística tradicional para la evaluación cuantitativa '
  'de la amenaza por movimientos en masa.')

# ══════════════════════════════════════════════════════════════════
# DISPONIBILIDAD DE DATOS
# ══════════════════════════════════════════════════════════════════
h('Disponibilidad de datos', level=1)

p('Los datos y códigos utilizados en este trabajo están disponibles en GitHub: '
  'https://github.com/edieraristizabal/NHPP')

# ══════════════════════════════════════════════════════════════════
# CONTRIBUCIÓN DE LOS AUTORES
# ══════════════════════════════════════════════════════════════════
h('Contribución de los autores', level=1)

p('[COMPLETAR: Describir la contribución del autor principal en el proceso de '
  'investigación, análisis de datos, escritura del manuscrito y revisión crítica.]')

# ══════════════════════════════════════════════════════════════════
# AGRADECIMIENTOS
# ══════════════════════════════════════════════════════════════════
h('Agradecimientos', level=1)

p('[COMPLETAR: Incluir únicamente personas que colaboraron en la realización del '
  'artículo o en la financiación del proyecto.]')

# ══════════════════════════════════════════════════════════════════
# REFERENCIAS (IEEE format)
# ══════════════════════════════════════════════════════════════════
h('Referencias', level=1)

refs = [
    '[1]\tP. Reichenbach, M. Rossi, B. D. Malamud, M. Mihir, and F. Guzzetti, '
    '"A review of statistically-based landslide susceptibility models," '
    'Earth-Science Reviews, vol. 180, pp. 60–91, 2018. '
    'https://doi.org/10.1016/j.earscirev.2018.03.001',

    '[2]\tF. Guzzetti, P. Reichenbach, F. Ardizzone, M. Cardinali, and M. Galli, '
    '"Estimating the quality of landslide susceptibility models," '
    'Geomorphology, vol. 81, no. 1–2, pp. 166–184, 2006. '
    'https://doi.org/10.1016/j.geomorph.2006.03.036',

    '[3]\tL. Lombardo and P. M. Mai, "Presenting logistic regression-based landslide '
    'susceptibility results," Engineering Geology, vol. 244, pp. 14–24, 2018. '
    'https://doi.org/10.1016/j.enggeo.2018.07.019',

    '[4]\tJ. Besag, "Spatial interaction and the statistical analysis of lattice '
    'systems," Journal of the Royal Statistical Society: Series B, vol. 36, no. 2, '
    'pp. 192–236, 1974. https://doi.org/10.1111/j.2517-6161.1974.tb00999.x',

    '[5]\tJ. Besag, J. York, and A. Mollie, "Bayesian image restoration, with two '
    'applications in spatial statistics," Annals of the Institute of Statistical '
    'Mathematics, vol. 43, no. 1, pp. 1–59, 1991. '
    'https://doi.org/10.1007/BF00116466',

    '[6]\tS. Openshaw, The Modifiable Areal Unit Problem. Norwich, England: '
    'Geo Books, 1984.',

    '[7]\tA. S. Fotheringham and D. W. S. Wong, "The modifiable areal unit problem '
    'in multivariate statistical analysis," Environment and Planning A, vol. 23, '
    'no. 7, pp. 1025–1044, 1991. https://doi.org/10.1068/a231025',

    '[8]\tP. J. Diggle, Statistical Analysis of Spatial and Spatio-Temporal Point '
    'Patterns, 3rd ed. Boca Raton, FL: CRC Press, 2013.',

    '[9]\tL. Lombardo, T. Opitz, and R. Huser, "Point process-based modeling of '
    'multiple debris flow landslides using INLA: An application to the 2009 Messina '
    'disaster," Stochastic Environmental Research and Risk Assessment, vol. 33, '
    'no. 7, pp. 1461–1478, 2019. https://doi.org/10.1007/s00477-019-01684-4',

    '[10]\tL. Lombardo, T. Opitz, F. Ardizzone, F. Guzzetti, and R. Huser, '
    '"Space-time landslide predictive modelling," Earth-Science Reviews, vol. 209, '
    'pp. 103318, 2020. https://doi.org/10.1016/j.earscirev.2020.103318',

    '[11]\tE. Aristizábal et al., "Susceptibilidad a los movimientos en masa mediante '
    'un modelo lineal generalizado jerárquico de Poisson con dependencia espacial '
    'sobre unidades de ladera: aplicación al departamento de Antioquia, Colombia," '
    'Andean Geology, 2025, en prensa.',

    '[12]\tE. V. Aristizábal, O. I. Sanchez, and O. Korup, "Landslide timing and '
    'rainfall regimes in a rapidly urbanizing tropical mountain valley in Colombia," '
    'Natural Hazards, vol. 122, pp. 306, 2026. '
    'https://doi.org/10.1007/s11069-026-08028-6',

    '[13]\tR. Trenkamp, J. N. Kellogg, J. T. Freymueller, and H. P. Mora, "Wide '
    'plate margin deformation, southern Central America and northwestern South '
    'America, CASA GPS observations," Journal of South American Earth Sciences, '
    'vol. 15, no. 2, pp. 157–171, 2002. '
    'https://doi.org/10.1016/S0895-9811(02)00018-4',

    '[14]\tF. Cediel, R. P. Shaw, and C. Cáceres, "Tectonic assembly of the northern '
    'Andean block," in The Circum-Gulf of Mexico and the Caribbean, C. Bartolini, '
    'R. T. Buffler, and J. Blickwede, Eds. American Association of Petroleum '
    'Geologists, 2003, pp. 815–848.',

    '[15]\tE. Aristizábal and O. Korup, "Linking landslide patterns to transient '
    'landscapes in the northern Colombian Andes," Journal of Geophysical Research: '
    'Earth Surface, vol. 130, pp. e2024JF008027, 2025. '
    'https://doi.org/10.1029/2024JF008027',

    '[16]\tD. Gómez, E. F. García, and E. Aristizábal, "Spatial and temporal '
    'landslide distributions using global and open landslide databases," '
    'Natural Hazards, vol. 117, no. 1, pp. 25–55, 2023. '
    'https://doi.org/10.1007/s11069-023-05848-8',

    '[17]\tJ. M. Bedoya-Soto, E. Aristizábal, A. M. Carmona, and G. Poveda, '
    '"Seasonal shift of the diurnal cycle of rainfall over Medellín\'s Valley, '
    'Central Andes of Colombia (1998–2005)," Frontiers in Earth Science, vol. 7, '
    'pp. 92, 2019. https://doi.org/10.3389/feart.2019.00092',

    '[18]\tG. Poveda, D. M. Alvarez, and O. A. Rueda, "Hydro-climatic variability '
    'over the Andes of Colombia associated with ENSO: A review of climatic processes '
    'and their impact on one of the Earth\'s most important biodiversity hotspots," '
    'Climate Dynamics, vol. 36, no. 11–12, pp. 2233–2249, 2011. '
    'https://doi.org/10.1007/s00382-010-0931-y',

    '[19]\tE. Aristizábal and O. Sánchez, "Spatial and temporal patterns and the '
    'socioeconomic impacts of landslides in the tropical and mountainous Colombian '
    'Andes," Disasters, vol. 44, no. 3, pp. 596–618, 2020. '
    'https://doi.org/10.1111/disa.12391',

    '[20]\tE. Aristizábal, B. Roser, and S. Yokota, "Tropical chemical weathering '
    'of hillslope deposits and bedrock source in the Aburrá valley, northern '
    'Colombian Andes," Engineering Geology, vol. 81, no. 4, pp. 389–406, 2005. '
    'https://doi.org/10.1016/j.enggeo.2005.08.001',

    '[21]\tH. Rue, S. Martino, and N. Chopin, "Approximate Bayesian inference for '
    'latent Gaussian models by using integrated nested Laplace approximations," '
    'Journal of the Royal Statistical Society: Series B, vol. 71, no. 2, '
    'pp. 319–392, 2009. https://doi.org/10.1111/j.1467-9868.2008.00700.x',

    '[22]\tH. Rue et al., "Bayesian computing with INLA: A review," Annual Review '
    'of Statistics and Its Application, vol. 4, pp. 395–421, 2017. '
    'https://doi.org/10.1146/annurev-statistics-060116-054045',

    '[23]\tR Core Team, R: A Language and Environment for Statistical Computing. '
    'Vienna, Austria: R Foundation for Statistical Computing, 2024. [Online]. '
    'Available: https://www.R-project.org/',

    '[24]\tF. E. Bachl, F. Lindgren, D. L. Borchers, and J. B. Illian, "inlabru: '
    'An R package for Bayesian spatial modelling from ecological survey data," '
    'Methods in Ecology and Evolution, vol. 10, no. 6, pp. 760–766, 2019. '
    'https://doi.org/10.1111/2041-210X.13168',

    '[25]\tF. Guzzetti et al., "Landslide inventory maps: New tools for an old '
    'problem," Earth-Science Reviews, vol. 112, no. 1–2, pp. 42–66, 2012. '
    'https://doi.org/10.1016/j.earscirev.2012.02.001',
]

for ref in refs:
    ref_para = doc.add_paragraph(style='Normal')
    ref_para.paragraph_format.first_line_indent = Cm(-0.7)
    ref_para.paragraph_format.left_indent = Cm(0.7)
    ref_para.paragraph_format.space_after = Pt(0)
    ref_para.add_run(ref)

# ── Save as docx ────────────────────────────────────────────────
OUT_DIR = '/home/edier/Documents/INVESTIGACION/PAPERS/ELABORACION/NHPP-movimientos-masa-VdeA'
docx_path = os.path.join(OUT_DIR, 'manuscrito_journal.docx')
doc.save(docx_path)
print(f"DOCX saved: {docx_path}")
